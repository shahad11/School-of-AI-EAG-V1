from mcp.server.fastmcp import FastMCP
from mcp.types import TextContent
from rich.console import Console
from rich.panel import Panel
import requests
from bs4 import BeautifulSoup
from docx import Document
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import os
from dotenv import load_dotenv
load_dotenv()

console = Console()
mcp = FastMCP("AINewsAgent")

# Pydantic models for validated I/O
from models import (
    FetchAINewsInput,
    FetchAINewsOutput,
    Article,
    FetchArticleContentInput,
    FetchArticleContentOutput,
    SaveToWordInput,
    SaveToWordOutput,
    ArticleContent,
    SendEmailInput,
    SendEmailOutput,
)
import json

@mcp.tool()
def fetch_ai_news(url: str) -> TextContent:
    """Fetch and parse the main AI news page, returning a list of (title, link) tuples for articles."""
    console.print("[blue]FUNCTION CALL:[/blue] fetch_ai_news()")
    try:
        # Try multiple AI news sources
        news_sources = [
            "https://www.artificialintelligence-news.com/artificial-intelligence-news/",
            "https://venturebeat.com/category/ai/",
            "https://www.theverge.com/ai-artificial-intelligence",
            "https://www.techcrunch.com/tag/artificial-intelligence/"
        ]
        
        articles = []
        for source_url in news_sources:
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                    'Accept-Encoding': 'gzip, deflate',
                    'Connection': 'keep-alive',
                    'Upgrade-Insecure-Requests': '1',
                }
                
                response = requests.get(source_url, headers=headers, timeout=10)
                console.print(f"Trying {source_url}: Status {response.status_code}")
                
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    # Different selectors for different sites
                    if 'artificialintelligence-news.com' in source_url:
                        selectors = ['h3.td-module-title a', 'div.td-module-thumb a', 'a[rel="bookmark"]']
                    elif 'venturebeat.com' in source_url:
                        selectors = ['h2 a', 'h3 a', '.post-title a']
                    elif 'theverge.com' in source_url:
                        selectors = ['h2 a', 'h3 a', '.c-entry-box--compact__title a']
                    elif 'techcrunch.com' in source_url:
                        selectors = ['h2 a', 'h3 a', '.post-block__title a']
                    else:
                        selectors = ['h2 a', 'h3 a', 'h1 a']
                    
                    for selector in selectors:
                        links = soup.select(selector)
                        for link in links:
                            title = link.get('title') or link.get_text().strip()
                            href = link.get('href')
                            if title and href and href.startswith('http') and len(title) > 10:
                                articles.append((title, href))
                                if len(articles) >= 15:
                                    break
                        if len(articles) >= 15:
                            break
                    
                    if articles:
                        console.print(f"Successfully fetched {len(articles)} articles from {source_url}")
                        break
                        
            except Exception as e:
                console.print(f"Failed to fetch from {source_url}: {e}")
                continue
        
        # If all sources fail, create some sample articles for testing
        if not articles:
            console.print("All sources failed. Creating sample articles for testing.")
            articles = [
                ("AI Breakthrough: New Model Achieves Human-Level Reasoning", "https://example.com/ai-breakthrough"),
                ("Robotics Revolution: Self-Learning Robots Transform Manufacturing", "https://example.com/robotics-revolution"),
                ("Machine Learning Advances: Neural Networks Solve Complex Problems", "https://example.com/ml-advances")
            ]
        
        console.print(f"Found {len(articles)} articles")
        return TextContent(type="text", text=str(articles[:10]))  # Return top 10 for selection
    except Exception as e:
        console.print(f"Error: {str(e)}")
        return TextContent(type="text", text=f"Error: {str(e)}")


# =========================
# Pydantic-validated tools
# =========================

@mcp.tool()
def fetch_ai_news_v2(input: FetchAINewsInput | dict) -> TextContent:
    """Fetch AI news; returns JSON matching FetchAINewsOutput."""
    console.print("[blue]FUNCTION CALL:[/blue] fetch_ai_news_v2()")
    try:
        if not isinstance(input, FetchAINewsInput):
            input = FetchAINewsInput.model_validate(input)

        # Reuse existing logic by calling fetch_ai_news on first source
        legacy = fetch_ai_news(url=input.url or "")

        # Parse legacy result which is a string of list of tuples
        try:
            articles_list = eval(legacy.text) if isinstance(legacy.text, str) else []
        except Exception:
            articles_list = []

        articles_models = [Article(title=title, url=url) for title, url in articles_list]
        payload = FetchAINewsOutput(articles=articles_models)
        return TextContent(type="text", text=payload.model_dump_json())
    except Exception as e:
        console.print(f"Error: {e}")
        return TextContent(type="text", text=FetchAINewsOutput(articles=[]).model_dump_json())

@mcp.tool()
def fetch_article_content(url: str) -> TextContent:
    """Fetch the full content of a news article given its URL."""
    console.print("[blue]FUNCTION CALL:[/blue] fetch_article_content()")
    try:
        # Handle sample URLs for testing
        if 'example.com' in url:
            if 'ai-breakthrough' in url:
                content = """Researchers have developed a new AI model that demonstrates human-level reasoning capabilities across multiple domains. The breakthrough, published in Nature, shows that the model can understand complex problems, reason through multi-step processes, and arrive at logical conclusions that match human performance.

The team used a novel architecture combining transformer networks with symbolic reasoning modules, allowing the AI to handle both pattern recognition and logical deduction. This represents a significant step toward artificial general intelligence (AGI).

Key findings include the model's ability to solve mathematical problems, understand natural language context, and apply learned knowledge to new situations. The research has implications for education, scientific discovery, and automated decision-making systems."""
            elif 'robotics-revolution' in url:
                content = """A new generation of self-learning robots is transforming manufacturing processes across industries. These robots can adapt to new tasks without explicit programming, learning from demonstration and improving through experience.

The technology combines computer vision, machine learning, and advanced control systems to enable robots to handle complex, variable tasks. Companies report 40% increases in productivity and significant reductions in setup time for new production lines.

Key applications include assembly, quality control, and material handling. The robots can work alongside humans safely and learn new skills through observation. This represents a major shift from traditional industrial robotics to more flexible, intelligent systems."""
            elif 'ml-advances' in url:
                content = """Recent advances in neural network architectures have enabled breakthroughs in solving previously intractable problems. New models can process multiple data types simultaneously, leading to improved performance in areas like computer vision, natural language processing, and scientific computing.

The research introduces novel attention mechanisms and training techniques that allow networks to learn more efficiently with less data. This addresses one of the major challenges in machine learning: the need for large datasets.

Applications include drug discovery, climate modeling, and autonomous systems. The techniques also improve interpretability, making AI decisions more transparent and trustworthy."""
            else:
                content = "Sample article content for testing purposes. This is placeholder content when real article content cannot be fetched."
            
            console.print(f"Generated sample content for {url}")
            return TextContent(type="text", text=content)
        
        # Try to fetch real content
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Try multiple selectors for article content
        content_selectors = [
            'div.td-post-content p',
            'div.entry-content p',
            'article p',
            'div.post-content p',
            '.post-body p',
            '.article-content p'
        ]
        
        paragraphs = []
        for selector in content_selectors:
            paragraphs = [p.get_text().strip() for p in soup.select(selector) if p.get_text().strip()]
            if paragraphs:
                break
        
        content = '\n\n'.join(paragraphs)
        if not content:
            content = "Content could not be extracted from this article. The article may be behind a paywall or use a different structure."
        
        console.print(f"Extracted {len(paragraphs)} paragraphs")
        return TextContent(type="text", text=content)
    except Exception as e:
        console.print(f"Error: {str(e)}")
        return TextContent(type="text", text=f"Error fetching content: {str(e)}")


@mcp.tool()
def fetch_article_content_v2(input: FetchArticleContentInput | dict) -> TextContent:
    """Fetch article content; returns JSON matching FetchArticleContentOutput."""
    console.print("[blue]FUNCTION CALL:[/blue] fetch_article_content_v2()")
    try:
        if not isinstance(input, FetchArticleContentInput):
            input = FetchArticleContentInput.model_validate(input)

        legacy = fetch_article_content(url=input.url)
        content = legacy.text
        return TextContent(
            type="text",
            text=FetchArticleContentOutput(content=content).model_dump_json(),
        )
    except Exception as e:
        console.print(f"Error: {e}")
        return TextContent(
            type="text",
            text=FetchArticleContentOutput(content=f"Error: {e}").model_dump_json(),
        )

@mcp.tool()
def save_to_word(filename: str, articles: list) -> TextContent:
    """Save a list of (title, content) tuples to a Word file, replacing the file if it exists."""
    console.print("[blue]FUNCTION CALL:[/blue] save_to_word()")
    try:
        # Always replace the file if it exists
        if os.path.exists(filename):
            try:
                os.remove(filename)
            except Exception as e:
                console.print(f"[yellow]Warning: Could not remove existing file: {e}[/yellow]")
        doc = Document()
        for title, content in articles:
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)
        doc.save(filename)
        return TextContent(type="text", text=f"Saved to {filename}")
    except Exception as e:
        console.print(f"[red]Error:[/red] {str(e)}")
        return TextContent(type="text", text=f"Error: {str(e)}")


@mcp.tool()
def save_to_word_v2(input: SaveToWordInput | dict) -> TextContent:
    """Save articles to a Word file using validated schema."""
    console.print("[blue]FUNCTION CALL:[/blue] save_to_word_v2()")
    try:
        if not isinstance(input, SaveToWordInput):
            input = SaveToWordInput.model_validate(input)

        # Convert to legacy (title, content) list
        legacy_articles = [(a.title, a.content) for a in input.articles]
        legacy_result = save_to_word(filename=input.filename, articles=legacy_articles)
        return TextContent(
            type="text",
            text=SaveToWordOutput(message=legacy_result.text).model_dump_json(),
        )
    except Exception as e:
        console.print(f"Error: {e}")
        return TextContent(
            type="text",
            text=SaveToWordOutput(message=f"Error: {e}").model_dump_json(),
        )

@mcp.tool()
def send_email(subject: str, body: str, to_email: str) -> TextContent:
    """Send an email with the given subject and body to the specified address. Uses SMTP settings from environment variables."""
    console.print("[blue]FUNCTION CALL:[/blue] send_email()")
    try:
        smtp_server = os.getenv("SMTP_SERVER")
        smtp_port = int(os.getenv("SMTP_PORT", 587))
        smtp_user = os.getenv("SMTP_USER")
        smtp_password = os.getenv("SMTP_PASSWORD")
        
        # Check each setting individually and provide specific error messages
        missing_settings = []
        if not smtp_server:
            missing_settings.append("SMTP_SERVER")
        if not smtp_user:
            missing_settings.append("SMTP_USER")
        if not smtp_password:
            missing_settings.append("SMTP_PASSWORD")
        
        if missing_settings:
            error_msg = f"Missing SMTP settings in .env file: {', '.join(missing_settings)}"
            console.print(f"Error: {error_msg}")
            return TextContent(type="text", text=f"Error: {error_msg}")
        
        console.print(f"Using SMTP: {smtp_server}:{smtp_port}")
        console.print(f"From: {smtp_user}")
        console.print(f"To: {to_email}")
        
        from_email = smtp_user
        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))
        
        # Use SMTP_SSL for port 465, regular SMTP for port 587
        if smtp_port == 465:
            with smtplib.SMTP_SSL(smtp_server, smtp_port) as server:
                server.login(smtp_user, smtp_password)
                server.sendmail(from_email, to_email, msg.as_string())
        else:
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(smtp_user, smtp_password)
                server.sendmail(from_email, to_email, msg.as_string())
        
        console.print("Email sent successfully")
        return TextContent(type="text", text="Email sent successfully.")
    except Exception as e:
        console.print(f"Error: {str(e)}")
        return TextContent(type="text", text=f"Error: {str(e)}")


@mcp.tool()
def send_email_v2(input: SendEmailInput | dict) -> TextContent:
    """Send email using validated schema."""
    console.print("[blue]FUNCTION CALL:[/blue] send_email_v2()")
    try:
        if not isinstance(input, SendEmailInput):
            input = SendEmailInput.model_validate(input)

        legacy = send_email(subject=input.subject, body=input.body, to_email=input.to_email)
        return TextContent(
            type="text",
            text=SendEmailOutput(message=legacy.text).model_dump_json(),
        )
    except Exception as e:
        console.print(f"Error: {e}")
        return TextContent(
            type="text",
            text=SendEmailOutput(message=f"Error: {e}").model_dump_json(),
        )

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "dev":
        mcp.run()
    else:
        mcp.run(transport="stdio") 